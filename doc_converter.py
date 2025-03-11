from docx import Document
from PyPDF2 import PdfReader, PdfWriter
import os
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import tkinter as tk
from tkinter import filedialog
import pdfplumber

# 注册支持中文的字体，这里假设系统中存在 SimSun.ttf 字体文件
pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))

class DocConverter:
    def __init__(self):
        pass

    def word_to_pdf(self, word_path, pdf_path):
        try:
            doc = Document(word_path)
            pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            # 设置支持中文的字体
            styles['Normal'].fontName = 'SimSun'
            story = []

            for para in doc.paragraphs:
                text = para.text
                style = styles['Normal']
                story.append(Paragraph(text, style))

            pdf.build(story)
            print(f"成功将 {word_path} 转换为 {pdf_path}")
            return True
        except Exception as e:
            print(f"转换过程中出现错误: {e}")
            return False

    def pdf_to_word(self, pdf_path, word_path):
        try:
            doc = Document()
            # 使用 pdfplumber 提取文本
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)

            doc.save(word_path)
            print(f"成功将 {pdf_path} 转换为 {word_path}")
            return True
        except Exception as e:
            print(f"转换过程中出现错误: {e}")
            return False

def select_input_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx"), ("PDF Files", "*.pdf")])
    input_entry.delete(0, tk.END)
    input_entry.insert(0, file_path)

def select_output_file():
    output_format = var.get()
    if output_format == "pdf":
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
    else:
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
    output_entry.delete(0, tk.END)
    output_entry.insert(0, file_path)

def convert_file():
    input_file = input_entry.get()
    output_file = output_entry.get()
    output_format = var.get()

    converter = DocConverter()
    if output_format == "pdf":
        success = converter.word_to_pdf(input_file, output_file)
    else:
        success = converter.pdf_to_word(input_file, output_file)

    if success:
        result_label.config(text="转换成功！")
    else:
        result_label.config(text="转换失败，请检查文件路径和格式。")

# 创建主窗口
root = tk.Tk()
root.title("文档格式转换工具")

# 输入文件选择
input_label = tk.Label(root, text="输入文件:")
input_label.pack()
input_entry = tk.Entry(root, width=50)
input_entry.pack()
input_button = tk.Button(root, text="选择文件", command=select_input_file)
input_button.pack()

# 输出文件选择
output_label = tk.Label(root, text="输出文件:")
output_label.pack()
output_entry = tk.Entry(root, width=50)
output_entry.pack()
output_button = tk.Button(root, text="选择保存路径", command=select_output_file)
output_button.pack()

# 转换目标格式选择
var = tk.StringVar()
var.set("pdf")
pdf_radio = tk.Radiobutton(root, text="转换为 PDF", variable=var, value="pdf")
pdf_radio.pack()
docx_radio = tk.Radiobutton(root, text="转换为 Word", variable=var, value="docx")
docx_radio.pack()

# 转换按钮
convert_button = tk.Button(root, text="开始转换", command=convert_file)
convert_button.pack()

# 结果显示
result_label = tk.Label(root, text="")
result_label.pack()

# 运行主循环
root.mainloop()